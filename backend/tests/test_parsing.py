import io

import pytest
from docx import Document

from app.services.parsing import (
    TextDecodeError,
    UnsupportedFileType,
    detect_content_type,
    extract_text,
)


def minimal_pdf(text: str) -> bytes:
    """xref 테이블 오프셋을 계산해 넣은 최소 PDF. text는 ASCII만 가능하다."""
    stream = f"BT /F1 24 Tf 72 700 Td ({text}) Tj ET".encode()
    objs = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        (
            b"<</Type/Page/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>"
            b"/MediaBox[0 0 612 792]/Contents 5 0 R>>"
        ),
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        b"<</Length %d>>\nstream\n%s\nendstream" % (len(stream), stream),
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objs, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj " % i + obj + b" endobj\n"
    xref_at = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objs) + 1)
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset
    out += b"trailer <</Size %d/Root 1 0 R>>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objs) + 1,
        xref_at,
    )
    return bytes(out)


@pytest.mark.parametrize(
    ("filename", "expected"),
    [
        ("report.pdf", "pdf"),
        ("manual.DOCX", "docx"),
        ("notes.txt", "txt"),
        ("README.md", "md"),
    ],
)
def test_detect_content_type(filename: str, expected: str) -> None:
    assert detect_content_type(filename) == expected


@pytest.mark.parametrize("filename", ["README", "plan.hwp", "data.xlsx"])
def test_detect_content_type_rejects_unsupported_files(filename: str) -> None:
    with pytest.raises(UnsupportedFileType):
        detect_content_type(filename)


@pytest.mark.parametrize("content_type", ["txt", "md"])
def test_extract_text_decodes_utf8_and_preserves_paragraph_breaks(content_type: str) -> None:
    text = "첫 번째 문단\n\n두 번째 문단"

    assert extract_text(text.encode(), content_type) == text


def test_extract_text_rejects_non_utf8_plain_text() -> None:
    with pytest.raises(TextDecodeError):
        extract_text("한국어".encode("cp949"), "txt")


def test_extract_text_reads_docx_paragraphs_in_order() -> None:
    buf = io.BytesIO()
    document = Document()
    document.add_paragraph("임베딩 잡은 트리거가 만든다")
    document.add_paragraph("두 번째 문단")
    document.save(buf)

    assert extract_text(buf.getvalue(), "docx") == (
        "임베딩 잡은 트리거가 만든다\n\n두 번째 문단"
    )


def test_extract_text_reads_pdf() -> None:
    assert extract_text(minimal_pdf("embedding job trigger"), "pdf") == "embedding job trigger"


def test_extract_text_returns_empty_string_for_pdf_without_text() -> None:
    assert extract_text(minimal_pdf(""), "pdf") == ""


@pytest.mark.parametrize("content_type", ["pdf", "docx"])
@pytest.mark.parametrize("data", [b"", b"not a document"])
def test_extract_text_normalizes_invalid_binary_document_errors(
    data: bytes, content_type: str
) -> None:
    with pytest.raises(ValueError, match="파일을 읽을 수 없습니다"):
        extract_text(data, content_type)


def test_extract_text_rejects_unsupported_content_type() -> None:
    with pytest.raises(UnsupportedFileType):
        extract_text(b"data", "hwp")
